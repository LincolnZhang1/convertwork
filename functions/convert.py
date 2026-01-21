import base64
from PIL import Image
import docx
from docx.shared import Inches
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import moviepy.editor as mp
from pydub import AudioSegment
import cv2
import requests
from bs4 import BeautifulSoup
import yt_dlp

# 简化版 - 使用标准库
app = None

# 上传目录 - 使用/tmp在Netlify Functions中
UPLOAD_DIR = Path("/tmp/uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

def get_file_type(filename: str) -> str:
    ext = filename.lower().split('.')[-1] if '.' in filename else ''
    
    image_exts = ['jpg', 'jpeg', 'png', 'gif', 'webp', 'bmp']
    document_exts = ['pdf', 'doc', 'docx', 'txt', 'html']
    audio_exts = ['mp3', 'wav', 'flac', 'aac', 'ogg', 'm4a']
    video_exts = ['mp4', 'avi', 'mkv', 'mov', 'webm', 'flv']
    
    if ext in image_exts:
        return 'image'
    elif ext in document_exts:
        return 'document'
    elif ext in audio_exts:
        return 'audio'
    elif ext in video_exts:
        return 'video'
    return 'other'

def get_mime_type(filename: str) -> str:
    mime_type, _ = mimetypes.guess_type(filename)
    return mime_type or 'application/octet-stream'

def handler(event, context):
    """Netlify函数入口"""
    
    try:
        # 解析请求
        if event.get('httpMethod') == 'GET':
            return handle_get(event)
        elif event.get('httpMethod') == 'POST':
            return handle_post(event)
        else:
            return {
                'statusCode': 405,
                'headers': {
                    'Content-Type': 'application/json',
                    'Access-Control-Allow-Origin': '*',
                    'Access-Control-Allow-Headers': 'Content-Type',
                },
                'body': json.dumps({'error': 'Method not allowed'})
            }
    
    except Exception as e:
        return {
            'statusCode': 500,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Headers': 'Content-Type',
            },
            'body': json.dumps({'error': f'Server error: {str(e)}'})
        }
    
    finally:
        # 清理临时文件
        try:
            if UPLOAD_DIR.exists():
                shutil.rmtree(UPLOAD_DIR)
        except:
            pass

def handle_get(event):
    """处理GET请求"""
    
    path = event.get('path', '')
    
    if path == '/api/status':
        # 检查系统依赖
        status = {
            "libreoffice": shutil.which('libreoffice') is not None,
            "ffmpeg": shutil.which('ffmpeg') is not None,
            "yt_dlp": shutil.which('yt-dlp') is not None,
            "python_imaging": True,
            "pypdf2": True,
        }
        
        return {
            'statusCode': 200,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*',
            },
            'body': json.dumps(status)
        }
    
    elif path.startswith('/api/download/'):
        # 文件下载 - 直接返回文件内容
        filename = path.split('/api/download/')[-1]
        file_path = UPLOAD_DIR / filename
        
        if not file_path.exists():
            return {
                'statusCode': 404,
                'headers': {
                    'Content-Type': 'application/json',
                    'Access-Control-Allow-Origin': '*',
                },
                'body': json.dumps({'error': 'File not found'})
            }
        
        # 读取文件
        with open(file_path, 'rb') as f:
            file_content = f.read()
        
        mime_type = get_mime_type(filename)
        
        return {
            'statusCode': 200,
            'headers': {
                'Content-Type': mime_type,
                'Content-Disposition': f'attachment; filename="{filename}"',
                'Access-Control-Allow-Origin': '*',
                'Cache-Control': 'no-cache',
            },
            'body': base64.b64encode(file_content).decode('utf-8'),
            'isBase64Encoded': True
        }
    
    else:
        return {
            'statusCode': 404,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*',
            },
            'body': json.dumps({'error': 'Not found'})
        }

def handle_post(event):
    """处理POST请求"""
    
    path = event.get('path', '')
    
    if path == '/api/convert':
        return handle_convert(event)
    
    else:
        return {
            'statusCode': 404,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Headers': 'Content-Type',
            },
            'body': json.dumps({'error': 'Not found'})
        }

def handle_convert(event):
    """处理文件转换"""
    
    try:
        # 解析multipart/form-data
        body = event.get('body', '')
        is_base64 = event.get('isBase64Encoded', False)
        
        if is_base64:
            body = base64.b64decode(body)
        else:
            body = body.encode() if isinstance(body, str) else body
        
        content_type = event.get('headers', {}).get('content-type', '')
        
        if not body or 'multipart/form-data' not in content_type:
            return {
                'statusCode': 400,
                'headers': {
                    'Content-Type': 'application/json',
                    'Access-Control-Allow-Origin': '*',
                },
                'body': json.dumps({'error': 'Invalid content type'})
            }
        
        # 简单的multipart解析 - 查找boundary
        boundary = None
        for part in content_type.split(';'):
            part = part.strip()
            if part.startswith('boundary='):
                boundary = part.split('=', 1)[1].strip('"')
                break
        
        if not boundary:
            return {
                'statusCode': 400,
                'headers': {
                    'Content-Type': 'application/json',
                    'Access-Control-Allow-Origin': '*',
                },
                'body': json.dumps({'error': 'No boundary found'})
            }
        
        # 解析parts
        parts = body.split(f'--{boundary}'.encode())
        file_data = None
        conversion_type = 'document'
        operation = 'convert'
        target_format = 'docx'
        video_url = None
        webpage_url = None
        
        for part in parts:
            if b'Content-Disposition' in part:
                lines = part.split(b'\r\n')
                headers = {}
                body_start = 0
                
                for i, line in enumerate(lines):
                    if line.startswith(b'Content-Disposition'):
                        # 解析Content-Disposition
                        disp = line.decode()
                        if 'name="file"' in disp:
                            # 文件部分
                            for j in range(i+1, len(lines)):
                                if lines[j].strip() == b'':
                                    body_start = j+1
                                    break
                            file_data = b'\r\n'.join(lines[body_start:-1])  # 去掉最后的\r\n
                        elif 'name="conversionType"' in disp:
                            for j in range(i+1, len(lines)):
                                if lines[j].strip() == b'':
                                    body_start = j+1
                                    break
                            conversion_type = b'\r\n'.join(lines[body_start:-1]).decode().strip()
                        elif 'name="operation"' in disp:
                            for j in range(i+1, len(lines)):
                                if lines[j].strip() == b'':
                                    body_start = j+1
                                    break
                            operation = b'\r\n'.join(lines[body_start:-1]).decode().strip()
                        elif 'name="targetFormat"' in disp:
                            for j in range(i+1, len(lines)):
                                if lines[j].strip() == b'':
                                    body_start = j+1
                                    break
                            target_format = b'\r\n'.join(lines[body_start:-1]).decode().strip()
                        elif 'name="videoUrl"' in disp:
                            for j in range(i+1, len(lines)):
                                if lines[j].strip() == b'':
                                    body_start = j+1
                                    break
                            video_url = b'\r\n'.join(lines[body_start:-1]).decode().strip()
                        elif 'name="webpageUrl"' in disp:
                            for j in range(i+1, len(lines)):
                                if lines[j].strip() == b'':
                                    body_start = j+1
                                    break
                            webpage_url = b'\r\n'.join(lines[body_start:-1]).decode().strip()
        
        # 生成唯一文件名
        file_id = str(uuid.uuid4())
        
        # 处理转换
        conversion_params = {
            'operation': operation,
            'targetFormat': target_format,
            'conversionType': conversion_type
        }
        
        if video_url:
            conversion_params['videoUrl'] = video_url
        if webpage_url:
            conversion_params['webpageUrl'] = webpage_url
        
        # 对于需要文件的操作
        if operation in ['convert'] and not file_data:
            return {
                'statusCode': 400,
                'headers': {
                    'Content-Type': 'application/json',
                    'Access-Control-Allow-Origin': '*',
                },
                'body': json.dumps({'error': 'No file uploaded'})
            }
        
        if file_data:
            original_filename = f"upload_{file_id}.bin"
            # 保存上传的文件
            file_path = UPLOAD_DIR / original_filename
            
            with open(file_path, 'wb') as f:
                f.write(file_data)
            
            file_type = get_file_type(original_filename)
        else:
            file_path = None
            file_type = conversion_type
        
        # 同步调用转换函数
        result = handle_conversion_sync(
            file_path, 
            original_filename if file_data else f"url_{file_id}", 
            conversion_params, file_id, file_type
        )
        return result
    
    except Exception as e:
        return {
            'statusCode': 500,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*',
            },
            'body': json.dumps({'error': f'Conversion failed: {str(e)}'})
        }

def handle_conversion_sync(file_path: Path, original_filename: str, params: Dict, file_id: str, file_type: str):
    """同步处理转换"""
    
    operation = params.get('operation', 'convert')
    target_format = params.get('targetFormat', 'pdf')
    conversion_type = params.get('conversionType', file_type)
    
    # 处理不需要文件的操作
    if 'videoUrl' in params and operation == 'download':
        return handle_youtube_download_sync(params['videoUrl'], target_format, file_id)
    elif 'webpageUrl' in params and operation == 'url-to-markdown':
        return handle_url_to_markdown_sync(params['webpageUrl'], file_id)
    
    # 需要文件的操作
    if not file_path or not file_path.exists():
        return {
            'statusCode': 400,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*',
            },
            'body': json.dumps({'success': False, 'error': 'File required for this operation'})
        }
    
    if operation == 'convert' and target_format:
        if conversion_type == 'image':
            return handle_image_conversion_sync(file_path, original_filename, target_format, file_id)
        elif conversion_type == 'document':
            return handle_document_conversion_sync(file_path, original_filename, target_format, file_id)
        elif conversion_type in ['audio', 'video']:
            return handle_media_conversion_sync(file_path, original_filename, target_format, conversion_type, file_id)
    
    return {
        'statusCode': 400,
        'headers': {
            'Content-Type': 'application/json',
            'Access-Control-Allow-Origin': '*',
        },
        'body': json.dumps({'success': False, 'error': 'Unsupported conversion'})
    }

def handle_image_conversion_sync(file_path: Path, original_filename: str, target_format: str, file_id: str):
    """图像转换"""
    
    try:
        output_filename = f"converted_{file_id}.{target_format}"
        output_path = UPLOAD_DIR / output_filename
        
        # 打开图像
        with Image.open(file_path) as img:
            # 转换格式
            if target_format.lower() in ['jpg', 'jpeg']:
                # 转换为RGB模式（如果需要）
                if img.mode in ('RGBA', 'LA', 'P'):
                    img = img.convert('RGB')
                img.save(output_path, 'JPEG', quality=90)
            elif target_format.lower() == 'png':
                img.save(output_path, 'PNG')
            elif target_format.lower() == 'webp':
                img.save(output_path, 'WEBP', quality=90)
            elif target_format.lower() == 'bmp':
                img.save(output_path, 'BMP')
            elif target_format.lower() == 'tiff':
                img.save(output_path, 'TIFF')
            else:
                return {
                    'statusCode': 400,
                    'headers': {
                        'Content-Type': 'application/json',
                        'Access-Control-Allow-Origin': '*',
                    },
                    'body': json.dumps({'success': False, 'error': f'Unsupported image format: {target_format}'})
                }
        
        return {
            'statusCode': 200,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*',
            },
            'body': json.dumps({
                'success': True,
                'downloadUrl': f"/api/download/{output_filename}",
                'fileName': output_filename,
                'fileSize': output_path.stat().st_size,
                'message': 'Image conversion completed'
            })
        }
        
    except Exception as e:
        return {
            'statusCode': 500,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*',
            },
            'body': json.dumps({'success': False, 'error': f'Image conversion failed: {str(e)}'})
        }

def handle_document_conversion_sync(file_path: Path, original_filename: str, target_format: str, file_id: str):
    """文档转换"""
    
    try:
        output_filename = f"converted_{file_id}.{target_format}"
        output_path = UPLOAD_DIR / output_filename
        
        ext = file_path.suffix.lower()
        
        if ext == '.pdf':
            if target_format.lower() == 'docx':
                # PDF到Word - 使用PyPDF2读取文本，然后创建Word文档
                from PyPDF2 import PdfReader
                
                reader = PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                
                # 创建Word文档
                doc = docx.Document()
                doc.add_paragraph(text)
                doc.save(output_path)
                
            elif target_format.lower() == 'txt':
                from PyPDF2 import PdfReader
                
                reader = PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                    
        elif ext in ['.docx', '.doc']:
            if target_format.lower() == 'pdf':
                # Word到PDF - 读取Word文档并转换为PDF
                import mammoth
                
                with open(file_path, "rb") as docx_file:
                    result = mammoth.convert_to_html(docx_file)
                    html = result.value
                
                # 简单的HTML到PDF转换（简化版）
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.platypus import SimpleDocTemplate, Paragraph
                
                doc = SimpleDocTemplate(str(output_path), pagesize=letter)
                styles = getSampleStyleSheet()
                story = []
                
                # 简单地将HTML转换为纯文本
                text = html.replace('<p>', '').replace('</p>', '\n').replace('<br>', '\n')
                story.append(Paragraph(text, styles["Normal"]))
                
                doc.build(story)
                
            elif target_format.lower() == 'txt':
                import mammoth
                
                with open(file_path, "rb") as docx_file:
                    result = mammoth.convert_to_html(docx_file)
                    html = result.value
                
                # 提取纯文本
                text = html.replace('<p>', '').replace('</p>', '\n').replace('<br>', '\n')
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                    
        elif ext == '.txt':
            if target_format.lower() == 'pdf':
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.platypus import SimpleDocTemplate, Paragraph
                
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                
                doc = SimpleDocTemplate(str(output_path), pagesize=letter)
                styles = getSampleStyleSheet()
                story = [Paragraph(text, styles["Normal"])]
                doc.build(story)
                
            elif target_format.lower() == 'docx':
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                
                doc = docx.Document()
                doc.add_paragraph(text)
                doc.save(output_path)
        
        if output_path.exists():
            return {
                'statusCode': 200,
                'headers': {
                    'Content-Type': 'application/json',
                    'Access-Control-Allow-Origin': '*',
                },
                'body': json.dumps({
                    'success': True,
                    'downloadUrl': f"/api/download/{output_filename}",
                    'fileName': output_filename,
                    'fileSize': output_path.stat().st_size,
                    'message': 'Document conversion completed'
                })
            }
        else:
            return {
                'statusCode': 500,
                'headers': {
                    'Content-Type': 'application/json',
                    'Access-Control-Allow-Origin': '*',
                },
                'body': json.dumps({'success': False, 'error': 'Conversion failed - output file not created'})
            }
        
    except Exception as e:
        return {
            'statusCode': 500,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*',
            },
            'body': json.dumps({'success': False, 'error': f'Document conversion failed: {str(e)}'})
        }
        
        return {
            'success': False,
            'error': f'Document conversion failed: {stderr.decode() if stderr else "Unknown error"}'
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': f'Document conversion failed: {str(e)}'
        }

def handle_media_conversion_sync(file_path: Path, original_filename: str, target_format: str, conversion_type: str, file_id: str):
    """音视频转换"""
    
    try:
        output_filename = f"converted_{file_id}.{target_format}"
        output_path = UPLOAD_DIR / output_filename
        
        if conversion_type == 'video':
            # 使用moviepy进行视频转换
            clip = mp.VideoFileClip(str(file_path))
            
            if target_format.lower() == 'mp4':
                clip.write_videofile(str(output_path), codec='libx264', audio_codec='aac', verbose=False, logger=None)
            elif target_format.lower() == 'webm':
                clip.write_videofile(str(output_path), codec='libvpx', audio_codec='libvorbis', verbose=False, logger=None)
            elif target_format.lower() == 'avi':
                clip.write_videofile(str(output_path), codec='png', audio_codec='pcm_s16le', verbose=False, logger=None)
            else:
                clip.write_videofile(str(output_path), verbose=False, logger=None)
            
            clip.close()
            
        elif conversion_type == 'audio':
            # 使用pydub进行音频转换
            audio = AudioSegment.from_file(str(file_path))
            
            if target_format.lower() == 'mp3':
                audio.export(str(output_path), format='mp3', bitrate='192k')
            elif target_format.lower() == 'wav':
                audio.export(str(output_path), format='wav')
            elif target_format.lower() == 'flac':
                audio.export(str(output_path), format='flac')
            elif target_format.lower() == 'aac':
                audio.export(str(output_path), format='adts', bitrate='128k')
            elif target_format.lower() == 'ogg':
                audio.export(str(output_path), format='ogg', bitrate='192k')
            else:
                audio.export(str(output_path), format=target_format)
        
        if output_path.exists():
            return {
                'statusCode': 200,
                'headers': {
                    'Content-Type': 'application/json',
                    'Access-Control-Allow-Origin': '*',
                },
                'body': json.dumps({
                    'success': True,
                    'downloadUrl': f"/api/download/{output_filename}",
                    'fileName': output_filename,
                    'fileSize': output_path.stat().st_size,
                    'message': 'Media conversion completed'
                })
            }
        else:
            return {
                'statusCode': 500,
                'headers': {
                    'Content-Type': 'application/json',
                    'Access-Control-Allow-Origin': '*',
                },
                'body': json.dumps({'success': False, 'error': 'Conversion failed - output file not created'})
            }
        
    except Exception as e:
        return {
            'statusCode': 500,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*',
            },
            'body': json.dumps({'success': False, 'error': f'Media conversion failed: {str(e)}'})
        }

def handle_youtube_download_sync(video_url: str, target_format: str, file_id: str):
    """YouTube视频下载"""
    
    try:
        output_filename = f"youtube_download_{file_id}.{target_format}"
        output_path = UPLOAD_DIR / output_filename
        
        # yt-dlp配置
        ydl_opts = {
            'outtmpl': str(output_path),
            'format': 'best[ext=mp4]/best' if target_format == 'mp4' else 'bestaudio/best',
        }
        
        if target_format in ['mp3', 'm4a', 'wav']:
            ydl_opts.update({
                'format': 'bestaudio/best',
                'postprocessors': [{
                    'key': 'FFmpegExtractAudio',
                    'preferredcodec': target_format if target_format != 'm4a' else 'aac',
                    'preferredquality': '192',
                }],
            })
        
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(video_url, download=True)
            title = info.get('title', 'youtube_video')
        
        # 重命名文件
        final_filename = f"{title.replace('/', '_').replace('\\', '_')}_{file_id}.{target_format}"
        final_path = UPLOAD_DIR / final_filename
        
        if output_path.exists():
            output_path.rename(final_path)
            output_filename = final_filename
            output_path = final_path
        
        return {
            'statusCode': 200,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*',
            },
            'body': json.dumps({
                'success': True,
                'downloadUrl': f"/api/download/{output_filename}",
                'fileName': output_filename,
                'fileSize': output_path.stat().st_size,
                'message': 'YouTube download completed'
            })
        }
        
    except Exception as e:
        return {
            'statusCode': 500,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*',
            },
            'body': json.dumps({'success': False, 'error': f'YouTube download failed: {str(e)}'})
        }

def handle_url_to_markdown_sync(webpage_url: str, file_id: str):
    """URL转Markdown"""
    
    try:
        # 获取网页内容
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(webpage_url, headers=headers, timeout=30)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # 提取标题
        title = soup.title.string if soup.title else 'Webpage'
        
        # 移除脚本和样式
        for script in soup(["script", "style"]):
            script.decompose()
        
        # 提取主要内容
        content = soup.get_text()
        
        # 清理文本
        lines = (line.strip() for line in content.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        content = '\n'.join(chunk for chunk in chunks if chunk)
        
        # 创建Markdown内容
        markdown_content = f"# {title}\n\n"
        markdown_content += f"Source: {webpage_url}\n\n"
        markdown_content += "---\n\n"
        markdown_content += content
        
        # 保存文件
        output_filename = f"webpage_{file_id}.md"
        output_path = UPLOAD_DIR / output_filename
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        return {
            'statusCode': 200,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*',
            },
            'body': json.dumps({
                'success': True,
                'downloadUrl': f"/api/download/{output_filename}",
                'fileName': output_filename,
                'fileSize': output_path.stat().st_size,
                'message': 'URL to Markdown conversion completed'
            })
        }
        
    except Exception as e:
        return {
            'statusCode': 500,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*',
            },
            'body': json.dumps({'success': False, 'error': f'URL to Markdown conversion failed: {str(e)}'})
        }

# 本地测试入口
if __name__ == "__main__":
    # 简单测试
    test_event = {
        'httpMethod': 'POST',
        'path': '/api/convert',
        'headers': {'content-type': 'multipart/form-data'},
        'body': b'test file content'
    }
    
    result = lambda_handler(test_event, None)
    print(json.dumps(result, indent=2))